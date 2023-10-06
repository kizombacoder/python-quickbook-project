
from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()
#
document.add_picture('me.jpg',width=Inches(2.0) )
#
name = input('what is your name?')
speak('hello '+ name +'how are you today')
speak('what is your phone number')
phone_number = input('what is your phone number')
email= input('what is your email?')
document.add_paragraph(
    name + '|'+ phone_number+ '|' +email)
#about me
document.add_heading('About me')
about_me=input('tell about yourself?')
document.add_paragraph(about_me)
# work experiences
document.add_heading('Work Experiences')
p=document.add_paragraph()
company=input ('Enter company')
from_date=input('from date')
to_date=input('To Date')
p.add_run(company +' ').bold= True 
p.add_run(from_date +'- ' +to_date +'\n').italic=True


experience_details=input(
    'describe your expererience at the company' + company)
p.add_run(experience_details)
# more experiences
while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company')
        from_date = input('from date')
        to_date = input('To Date')
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '- ' + to_date + '\n').italic = True
        experience_details = input(
            'describe your experience at the company' + company +' ' )
        p.add_run(experience_details)
    else:
        break

#  skills

document.add_heading('skills')
p=document.add_paragraph()
skill=input ('what skills do you have?')
p.add_run(skill +' ' ).bold= True 

# more skills
while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No')
    if has_more_skills.lower() == 'yes':
        p = document.add_paragraph()
        skill = input('what skills do you have?')
        p = document.add_paragraph(skill)
        p.style = 'list bullet'
    else:
        break
#footer
section = document.sections [0]
footer= section.footer
p=footer.paragraphs[0]
p.text='cv generated by Dancecity'

document.save ('cv.docx')
